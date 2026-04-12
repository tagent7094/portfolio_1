"""Read all content from a docx file including tables, headers, etc."""
import sys
from docx import Document

path = sys.argv[1]
doc = Document(path)

# Paragraphs
paras = [p.text for p in doc.paragraphs if p.text.strip()]
print(f"=== PARAGRAPHS ({len(paras)}) ===")
for p in paras[:50]:
    print(p[:200])

# Tables
print(f"\n=== TABLES ({len(doc.tables)}) ===")
for ti, table in enumerate(doc.tables):
    print(f"\nTable {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip()[:100] for c in row.cells]
        print(f"  Row {ri}: {' | '.join(cells)}")
        if ri > 20:
            print(f"  ... ({len(table.rows) - 20} more rows)")
            break

# Sections / headers
print(f"\n=== SECTIONS ({len(doc.sections)}) ===")

# Raw XML element count
from docx.oxml.ns import qn
body = doc.element.body
children = list(body)
print(f"\n=== BODY ELEMENTS ({len(children)}) ===")
for c in children[:10]:
    print(f"  {c.tag}: {c.text[:100] if c.text else '(no text)'}")

# Total text extraction via runs
all_text = []
for para in doc.paragraphs:
    runs_text = ''.join(r.text for r in para.runs)
    if runs_text.strip():
        all_text.append(runs_text)

print(f"\n=== ALL RUNS TEXT ({len(all_text)} paragraphs) ===")
full = '\n'.join(all_text)
print(full[:5000])
print(f"\n=== TOTAL: {len(full)} chars ===")
