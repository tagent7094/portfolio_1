"""Read files from the founder-data directory."""

from __future__ import annotations

import csv
import json
import logging
import sys
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class Chunk:
    text: str
    source_file: str
    platform: str = "unknown"
    date: str | None = None
    engagement: int = 0
    metadata: dict = field(default_factory=dict)


def detect_platform(filename: str, text: str = "") -> str:
    """Guess which platform content came from based on filename or text."""
    name = filename.lower()
    for platform in ["linkedin", "twitter", "email", "blog", "podcast"]:
        if platform in name:
            return platform
    return "general"


def read_markdown(path: Path) -> list[Chunk]:
    """Read a markdown file and return chunks."""
    text = path.read_text(encoding="utf-8", errors="replace")
    return [Chunk(text=text, source_file=str(path), platform=detect_platform(path.name, text))]


def read_text(path: Path) -> list[Chunk]:
    """Read a plain text file and return chunks."""
    text = path.read_text(encoding="utf-8", errors="replace")
    return [Chunk(text=text, source_file=str(path), platform=detect_platform(path.name, text))]


def read_csv(path: Path) -> list[Chunk]:
    """Read a CSV file, expecting columns like content, date, likes."""
    import pandas as pd

    df = pd.read_csv(path)
    chunks = []
    # Try common column names for content
    content_col = None
    for col in ["content", "text", "body", "post", "message"]:
        if col in df.columns:
            content_col = col
            break
    if content_col is None:
        content_col = df.columns[0]

    date_col = next((c for c in df.columns if c.lower() in ("date", "created_at", "timestamp")), None)
    engagement_col = next((c for c in df.columns if c.lower() in ("likes", "engagement", "impressions")), None)
    platform_col = next((c for c in df.columns if c.lower() in ("platform", "source", "channel")), None)

    for _, row in df.iterrows():
        text = str(row[content_col])
        if not text or text == "nan":
            continue
        row_platform = str(row[platform_col]).strip() if platform_col and str(row[platform_col]) != "nan" else None
        chunks.append(
            Chunk(
                text=text,
                source_file=str(path),
                platform=row_platform or detect_platform(path.name, text),
                date=str(row[date_col]) if date_col and str(row[date_col]) != "nan" else None,
                engagement=int(row[engagement_col]) if engagement_col and str(row[engagement_col]) != "nan" else 0,
            )
        )
    return chunks


def read_docx(path: Path) -> list[Chunk]:
    """Read a .docx file and return chunks.

    Handles both paragraph-based and table-based content.
    For tables: looks for a content column (post/content/text) and optional engagement columns.
    """
    try:
        from docx import Document
    except ImportError:
        print(f"\033[35m[FileReader]\033[0m \033[31m-> Install python-docx: pip install python-docx\033[0m", file=sys.stderr, flush=True)
        logger.error("python-docx not installed. Run: pip install python-docx")
        return []

    doc = Document(str(path))
    chunks = []

    # Try paragraphs first
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        text = "\n\n".join(paragraphs)
        chunks.append(Chunk(text=text, source_file=str(path), platform=detect_platform(path.name, text)))

    # Also extract from tables (many founder data files use tables with post/likes/comments)
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        # Get header row
        headers = [c.text.strip().lower() for c in table.rows[0].cells]

        # Find content column
        content_idx = None
        for i, h in enumerate(headers):
            if h in ("post", "content", "text", "body", "message"):
                content_idx = i
                break
        if content_idx is None:
            content_idx = 0  # fallback to first column

        # Find engagement columns
        likes_idx = next((i for i, h in enumerate(headers) if "like" in h), None)
        comments_idx = next((i for i, h in enumerate(headers) if "comment" in h), None)
        reposts_idx = next((i for i, h in enumerate(headers) if "repost" in h or "share" in h), None)

        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            text = cells[content_idx] if content_idx < len(cells) else ""
            if not text or len(text) < 20:
                continue

            engagement = 0
            if likes_idx is not None and likes_idx < len(cells):
                try:
                    engagement += int(cells[likes_idx])
                except (ValueError, TypeError):
                    pass

            chunks.append(Chunk(
                text=text,
                source_file=str(path),
                platform=detect_platform(path.name, text),
                engagement=engagement,
            ))

    return chunks


def read_xlsx(path: Path) -> list[Chunk]:
    """Read an .xlsx file — treat like CSV with pandas."""
    import pandas as pd
    try:
        df = pd.read_excel(path)
    except Exception as e:
        logger.error("Failed to read xlsx %s: %s", path.name, e)
        return []

    # Same logic as read_csv
    content_col = None
    for col in ["content", "text", "body", "post", "message"]:
        if col in df.columns:
            content_col = col
            break
    if content_col is None:
        content_col = df.columns[0]

    chunks = []
    for _, row in df.iterrows():
        text = str(row[content_col])
        if not text or text == "nan":
            continue
        chunks.append(Chunk(text=text, source_file=str(path), platform=detect_platform(path.name, text)))
    return chunks


def read_json(path: Path) -> list[Chunk]:
    """Read a JSON file and extract text content."""
    data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    chunks = []
    if isinstance(data, list):
        for item in data:
            text = item.get("content") or item.get("text") or json.dumps(item)
            chunks.append(Chunk(text=text, source_file=str(path), platform=detect_platform(path.name)))
    elif isinstance(data, dict):
        text = data.get("content") or data.get("text") or json.dumps(data)
        chunks.append(Chunk(text=text, source_file=str(path), platform=detect_platform(path.name)))
    return chunks


READERS = {
    ".md": read_markdown,
    ".txt": read_text,
    ".csv": read_csv,
    ".json": read_json,
    ".docx": read_docx,
    ".xlsx": read_xlsx,
    ".xls": read_xlsx,
}


def read_all_files(directory: str = "data/founder-data") -> list[Chunk]:
    """Read all supported files from a directory."""
    print(f"\033[35m[FileReader]\033[0m \033[1mread_all_files(directory={directory!r})\033[0m", file=sys.stderr, flush=True)
    dir_path = Path(directory)
    if not dir_path.exists():
        print(f"\033[35m[FileReader]\033[0m \033[31m→ Directory does not exist!\033[0m", file=sys.stderr, flush=True)
        logger.warning("Directory %s does not exist", directory)
        return []

    chunks = []
    for path in sorted(dir_path.iterdir()):
        if path.name.startswith("."):
            continue
        reader = READERS.get(path.suffix.lower())
        if reader:
            print(f"\033[35m[FileReader]\033[0m Reading: {path.name} ({path.suffix})", file=sys.stderr, flush=True)
            logger.info("Reading %s", path.name)
            try:
                file_chunks = reader(path)
                print(f"\033[35m[FileReader]\033[0m   \033[32m→ {len(file_chunks)} chunks from {path.name}\033[0m", file=sys.stderr, flush=True)
                chunks.extend(file_chunks)
            except Exception as e:
                print(f"\033[35m[FileReader]\033[0m   \033[31m→ FAILED: {e}\033[0m", file=sys.stderr, flush=True)
                logger.error("Failed to read %s: %s", path.name, e)
        else:
            print(f"\033[35m[FileReader]\033[0m Skipping unsupported: {path.name}", file=sys.stderr, flush=True)
            logger.debug("Skipping unsupported file: %s", path.name)
    print(f"\033[35m[FileReader]\033[0m \033[32m→ Total: {len(chunks)} chunks from {directory}\033[0m", file=sys.stderr, flush=True)
    logger.info("Read %d chunks from %s", len(chunks), directory)
    return chunks
