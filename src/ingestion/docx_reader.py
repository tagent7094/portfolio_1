"""Extract text + tables from .docx files.

Promoted from scripts/read_docx_full.py into a proper module.
Handles password-protected and corrupt files gracefully.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import TypedDict

logger = logging.getLogger(__name__)


class DocxContent(TypedDict):
    paragraphs: list[str]
    tables: list[list[list[str]]]
    plain_text: str
    error: str


def read_docx(path: str | Path) -> DocxContent:
    """Read all content from a docx file.

    Returns a dict with `paragraphs`, `tables` (list of rows of cells), `plain_text`
    (paragraphs + table cells joined with newlines), and `error` (empty string on
    success; populated with a message on failure).

    Never raises — failures populate the `error` field so the caller can log
    and continue. This lets the universal reader keep walking the folder.
    """
    result: DocxContent = {
        "paragraphs": [],
        "tables": [],
        "plain_text": "",
        "error": "",
    }
    p = Path(path)
    if not p.exists():
        result["error"] = f"file not found: {p}"
        return result

    try:
        from docx import Document  # type: ignore
    except ImportError as e:
        result["error"] = f"python-docx not installed: {e}"
        logger.error("[docx_reader] %s", result["error"])
        return result

    try:
        doc = Document(str(p))
    except Exception as e:
        result["error"] = f"failed to open docx: {type(e).__name__}: {e}"
        logger.warning("[docx_reader] %s — %s", p.name, result["error"])
        return result

    try:
        paragraphs = [para.text for para in doc.paragraphs if para.text and para.text.strip()]
        result["paragraphs"] = paragraphs

        tables: list[list[list[str]]] = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append(rows)
        result["tables"] = tables

        text_parts = list(paragraphs)
        for table in tables:
            for row in table:
                non_empty = [c for c in row if c]
                if non_empty:
                    text_parts.append(" | ".join(non_empty))
        result["plain_text"] = "\n".join(text_parts)
    except Exception as e:
        result["error"] = f"failed to extract content: {type(e).__name__}: {e}"
        logger.warning("[docx_reader] %s — %s", p.name, result["error"])
        return result

    return result
